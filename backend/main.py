# from fastapi import FastAPI
# from pydantic import BaseModel
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.schema import Document
# from langchain_community.vectorstores import Qdrant
# from langchain.agents import Tool, initialize_agent
# from qdrant_client import QdrantClient
# from qdrant_client.http.models import Distance, VectorParams
# from langchain.embeddings.base import Embeddings
# import os
# from dotenv import load_dotenv
# import asyncio

# # ------------------ Setup ------------------
# load_dotenv()
# GOOGLE_API_KEY = os.getenv("GEMINI_API_KEY")
# QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
# QDRANT_URL = os.getenv("QDRANT_URL")  # Must be your Qdrant Cloud URL

# app = FastAPI()

# # ------------------ Text splitting ------------------
# text_splitter = RecursiveCharacterTextSplitter(
#     chunk_size=800,
#     chunk_overlap=100,
#     add_start_index=True
# )

# VECTOR_DIM = 768
# COLLECTION_NAME = "research_docs"

# # ------------------ Dummy embeddings ------------------
# class DummyEmbeddings(Embeddings):
#     def embed_documents(self, texts):
#         return [[0.0]*VECTOR_DIM for _ in texts]
#     def embed_query(self, text):
#         return [0.0]*VECTOR_DIM

# embeddings = DummyEmbeddings()

# # ------------------ Qdrant setup ------------------
# try:
#     qdrant_client = QdrantClient(
#         url=QDRANT_URL,
#         api_key=QDRANT_API_KEY,
#         timeout=60
#     )
#     qdrant_client.get_collections()
#     print("✅ Connected to Qdrant Cloud")
# except Exception as e:
#     qdrant_client = None
#     print(f"❌ Failed to connect to Qdrant: {e}")

# def ensure_collection():
#     """Ensure the collection exists in Qdrant Cloud"""
#     if qdrant_client is None:
#         return False
#     try:
#         qdrant_client.get_collection(COLLECTION_NAME)
#         return True
#     except Exception:
#         try:
#             qdrant_client.create_collection(
#                 collection_name=COLLECTION_NAME,
#                 vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
#             )
#             return True
#         except Exception as e:
#             print(f"❌ Could not create collection: {e}")
#             return False

# vector_store = None

# def get_or_create_vector_store():
#     global vector_store
#     if vector_store is None:
#         if qdrant_client is None:
#             raise RuntimeError("Qdrant is not available. Check your QDRANT_URL and API key.")
#         if ensure_collection():
#             vector_store = Qdrant(
#                 client=qdrant_client,
#                 collection_name=COLLECTION_NAME,
#                 embeddings=embeddings
#             )
#         else:
#             raise RuntimeError("Failed to initialize Qdrant collection.")
#     return vector_store

# # ------------------ Dummy Gemini LLM ------------------
# def get_gemini():
#     # For testing, we still call Gemini to generate answers, but embeddings are dummy
#     return ChatGoogleGenerativeAI(
#         model="gemini-2.0-flash-lite",
#         temperature=0.2,
#         api_key=GOOGLE_API_KEY
#     )

# # ------------------ Helpers ------------------
# def process_text(text: str):
#     """Split text into chunks and store in Qdrant"""
#     vs = get_or_create_vector_store()
#     doc = Document(page_content=text)
#     chunks = text_splitter.split_documents([doc])
#     try:
#         vs.add_documents(chunks)
#     except Exception as e:
#         print(f"⚠️ Failed to add documents to Qdrant: {e}")
#     return chunks

# # ------------------ Custom tools ------------------
# def research_tool(text: str):
#     """Store user-provided research text into knowledge base"""
#     try:
#         chunks = process_text(text)
#         return f"Stored research text as {len(chunks)} chunks in vector DB."
#     except Exception as e:
#         return f"Error storing research text: {e}"

# research_tool = Tool(
#     name="ResearchTool",
#     func=research_tool,
#     description="Stores research-related text in the knowledge base for retrieval."
# )

# # ------------------ Agent setup ------------------
# agent = initialize_agent(
#     tools=[research_tool],
#     llm=get_gemini(),
#     agent="zero-shot-react-description",
#     verbose=True
# )

# # ------------------ Request model ------------------
# class Query(BaseModel):
#     text: str

# @app.post("/chat")
# async def chat(query: Query):
#     user_text = query.text.strip()
#     print("Received query:", user_text)

#     loop = asyncio.get_event_loop()

#     # Retrieve relevant content from Qdrant
#     try:
#         vs = get_or_create_vector_store()
#         relevant_docs = await loop.run_in_executor(
#             None, lambda: vs.similarity_search(user_text, k=5)
#         )
#     except Exception as e:
#         return {"result": f"⚠️ Retrieval error: {str(e)}"}

#     # Build context for the LLM
#     context_text = ""
#     if relevant_docs:
#         context_text += "Here are relevant excerpts from knowledge base:\n"
#         for i, doc in enumerate(relevant_docs):
#             context_text += f"Chunk {i+1}:\n{doc.page_content}\n\n"

#     full_prompt = f"{context_text}\nUser query: {user_text}\nAnswer in detail using only the given excerpts. Cite sources if possible."

#     # Generate response
#     llm = get_gemini()
#     try:
#         result = await loop.run_in_executor(None, lambda: llm.invoke(full_prompt))
#         result_text = result.content if hasattr(result, "content") else str(result)
#         return {"result": result_text}
#     except Exception as e:
#         return {"result": f"⚠️ LLM error: {str(e)}"}











# from fastapi import FastAPI, UploadFile, File
# from pydantic import BaseModel
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.schema import Document
# from langchain_community.vectorstores import Qdrant
# from qdrant_client import QdrantClient
# from qdrant_client.http.models import Distance, VectorParams
# import os
# from dotenv import load_dotenv
# import asyncio
# import PyPDF2
# import docx
# from pdf2image import convert_from_bytes
# import pytesseract


# # ------------------ Setup ------------------
# load_dotenv()
# GOOGLE_API_KEY = os.getenv("GEMINI_API_KEY")
# QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
# QDRANT_URL = os.getenv("QDRANT_URL")

# app = FastAPI()

# VECTOR_DIM = 768
# COLLECTION_NAME = "uploaded_docs"
# file_uploaded = False   # ✅ Track whether a file has been uploaded

# # ------------------ Dummy embeddings ------------------
# from langchain.embeddings.base import Embeddings

# class DummyEmbeddings(Embeddings):
#     def embed_documents(self, texts):
#         return [[0.0]*VECTOR_DIM for _ in texts]
#     def embed_query(self, text):
#         return [0.0]*VECTOR_DIM

# embeddings = DummyEmbeddings()

# # ------------------ Qdrant setup ------------------
# try:
#     qdrant_client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=60)
#     collections = [c.name for c in qdrant_client.get_collections().collections]
#     if COLLECTION_NAME not in collections:
#         qdrant_client.create_collection(
#             collection_name=COLLECTION_NAME,
#             vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
#         )
#     vector_store = Qdrant(client=qdrant_client, collection_name=COLLECTION_NAME, embeddings=embeddings)
#     print("✅ Qdrant ready")
# except Exception as e:
#     vector_store = None
#     print(f"❌ Qdrant unavailable: {e}")

# text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100, add_start_index=True)

# # ------------------ LLM ------------------
# def get_gemini():
#     return ChatGoogleGenerativeAI(model="gemini-2.0-flash-lite", temperature=0.2, api_key=GOOGLE_API_KEY)

# # ------------------ Helpers ------------------
# def read_file(file: UploadFile) -> str:
#     """Read PDF, DOCX, or TXT files and return plain text."""
#     if file.filename.endswith(".pdf"):
#         reader = PyPDF2.PdfReader(file.file)
#         text = ""
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#         if not text.strip():
#             raise ValueError("❌ Could not extract text from PDF (might be image-based).")
#         return text

#     elif file.filename.endswith(".docx"):
#         doc = docx.Document(file.file)
#         return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

#     else:  # txt
#         return file.file.read().decode("utf-8")


# def process_text(text: str):
#     """Split text into chunks and store in Qdrant."""
#     if vector_store is None:
#         return []

#     # Clear old collection before inserting new file content
#     qdrant_client.delete_collection(collection_name=COLLECTION_NAME)
#     qdrant_client.create_collection(
#         collection_name=COLLECTION_NAME,
#         vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
#     )

#     doc = Document(page_content=text)
#     chunks = text_splitter.split_documents([doc])
#     vector_store.add_documents(chunks)
#     return chunks

# # ------------------ Request model ------------------
# class Query(BaseModel):
#     text: str

# # ------------------ Routes ------------------
# @app.post("/upload_file")
# async def upload_file(file: UploadFile = File(...)):
#     global file_uploaded
#     try:
#         content = read_file(file)
#         chunks = process_text(content)
#         file_uploaded = True   # ✅ Mark file as uploaded
#         return {"message": f"✅ File processed successfully with {len(chunks)} chunks."}
#     except Exception as e:
#         return {"message": f"❌ Error processing file: {e}"}

# @app.post("/chat")
# async def chat(query: Query):
#     user_text = query.text.strip()
#     loop = asyncio.get_event_loop()
#     response_text = ""

#     # ✅ Only search Qdrant if a file has been uploaded
#     if vector_store is not None and file_uploaded:
#         try:
#             relevant_docs = await loop.run_in_executor(
#                 None, lambda: vector_store.similarity_search(user_text, k=5)
#             )
#             if relevant_docs:
#                 context_text = "Here are relevant excerpts from the uploaded file:\n"
#                 for i, doc in enumerate(relevant_docs):
#                     context_text += f"Chunk {i+1}:\n{doc.page_content}\n\n"
#                 full_prompt = f"{context_text}\nUser query: {user_text}\nAnswer using only the above excerpts."
#                 llm = get_gemini()
#                 result = await loop.run_in_executor(None, lambda: llm.invoke(full_prompt))
#                 response_text = result.content if hasattr(result, "content") else str(result)
#             else:
#                 response_text = "⚠️ I couldn’t find relevant info in the uploaded document."
#         except Exception as e:
#             response_text = f"❌ Error retrieving from Qdrant: {e}"

#     else:
#         # ✅ If no file uploaded → fallback to direct LLM response
#         llm = get_gemini()
#         try:
#             result = await loop.run_in_executor(None, lambda: llm.invoke(user_text))
#             response_text = result.content if hasattr(result, "content") else str(result)
#         except Exception as e:
#             response_text = f"❌ Error generating LLM response: {e}"

#     return {"result": response_text}








# from fastapi import FastAPI, UploadFile, File
# from pydantic import BaseModel
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.schema import Document
# from langchain_community.vectorstores import Qdrant
# from qdrant_client import QdrantClient
# from qdrant_client.http.models import Distance, VectorParams
# import os
# from dotenv import load_dotenv
# import asyncio
# import PyPDF2
# import docx
# from pdf2image import convert_from_bytes
# import pytesseract

# # ------------------ Load environment variables ------------------
# load_dotenv()
# GOOGLE_API_KEY = os.getenv("GEMINI_API_KEY")
# QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
# QDRANT_URL = os.getenv("QDRANT_URL")

# app = FastAPI()

# VECTOR_DIM = 768
# COLLECTION_NAME = "uploaded_docs"
# file_uploaded = False  # Track if a file has been uploaded

# # ------------------ Embeddings ------------------
# # ⚠️ For Google Generative AI, it's better to use service account ADC for embeddings
# # You can also use this dummy embeddings for testing without Google credentials
# embeddings = GoogleGenerativeAIEmbeddings(
#     model="models/embedding-001",
#     google_api_key=GOOGLE_API_KEY  # use `google_api_key` parameter, not `api_key`
# )

# # ------------------ Qdrant Cloud Setup ------------------
# try:
#     qdrant_client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=60)
#     existing_collections = [c.name for c in qdrant_client.get_collections().collections]
#     if COLLECTION_NAME not in existing_collections:
#         qdrant_client.create_collection(
#             collection_name=COLLECTION_NAME,
#             vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
#         )
#     vector_store = Qdrant(client=qdrant_client, collection_name=COLLECTION_NAME, embeddings=embeddings)
#     print("✅ Qdrant Cloud is ready")
# except Exception as e:
#     vector_store = None
#     print(f"❌ Qdrant unavailable: {e}")

# # ------------------ Text Splitter ------------------
# text_splitter = RecursiveCharacterTextSplitter(chunk_size=2500, chunk_overlap=300, add_start_index=True)

# # ------------------ LLM ------------------
# def get_gemini():
#     return ChatGoogleGenerativeAI(model="gemini-2.0-flash-lite", temperature=0.2, api_key=GOOGLE_API_KEY)

# # ------------------ Helper Functions ------------------
# def read_file(file: UploadFile) -> str:
#     """Read PDF, DOCX, or TXT and return plain text."""
#     if file.filename.endswith(".pdf"):
#         file.file.seek(0)
#         reader = PyPDF2.PdfReader(file.file)
#         text = ""
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#         if not text.strip():  # OCR fallback
#             file.file.seek(0)
#             images = convert_from_bytes(file.file.read())
#             for img in images:
#                 text += pytesseract.image_to_string(img) + "\n"
#         if not text.strip():
#             raise ValueError("❌ Could not extract text from PDF.")
#         return text

#     elif file.filename.endswith(".docx"):
#         file.file.seek(0)
#         doc = docx.Document(file.file)
#         return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

#     else:  # txt
#         return file.file.read().decode("utf-8")

# def process_text(text: str):
#     """Split text into chunks and store in Qdrant Cloud."""
#     if vector_store is None:
#         return []

#     # Clear old collection content for new upload
#     try:
#         qdrant_client.delete_collection(collection_name=COLLECTION_NAME)
#     except Exception:
#         pass  # ignore if collection doesn't exist
#     qdrant_client.create_collection(
#         collection_name=COLLECTION_NAME,
#         vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
#     )

#     doc = Document(page_content=text)
#     chunks = text_splitter.split_documents([doc])
#     vector_store.add_documents(chunks)
#     return chunks

# # ------------------ Request Models ------------------
# class Query(BaseModel):
#     text: str

# class SearchQuery(BaseModel):
#     title: str

# # ------------------ Routes ------------------
# @app.post("/upload_file")
# async def upload_file(file: UploadFile = File(...)):
#     global file_uploaded
#     try:
#         content = read_file(file)
#         chunks = process_text(content)
#         file_uploaded = True
#         return {"message": f"✅ File processed successfully with {len(chunks)} chunks."}
#     except Exception as e:
#         return {"message": f"❌ Error processing file: {e}"}

# @app.post("/chat")
# async def chat(query: Query):
#     user_text = query.text.strip()
#     loop = asyncio.get_event_loop()
#     response_text = ""

#     if vector_store is not None and file_uploaded:
#         # ✅ Retrieve from uploaded docs in Qdrant
#         try:
#             relevant_docs = await loop.run_in_executor(
#                 None, lambda: vector_store.similarity_search(user_text, k=5)
#             )
#             if relevant_docs:
#                 context_text = "Here are relevant excerpts from the uploaded document:\n"
#                 for i, doc in enumerate(relevant_docs):
#                     context_text += f"Chunk {i+1}:\n{doc.page_content}\n\n"
#                 full_prompt = f"{context_text}\nUser query: {user_text}\nAnswer using only the above excerpts."
#                 llm = get_gemini()
#                 result = await loop.run_in_executor(None, lambda: llm.invoke(full_prompt))
#                 response_text = result.content if hasattr(result, "content") else str(result)
#             else:
#                 response_text = "⚠️ No relevant info found in uploaded document."
#         except Exception as e:
#             response_text = f"❌ Error retrieving from Qdrant: {e}"

#     else:
#         # ✅ No upload → Research Assistant: fetch related articles/papers
#         try:
#             llm = get_gemini()
#             prompt = (
#                 f"Provide 5 relevant academic papers, research articles, or web sources "
#                 f"related to: '{user_text}'. Include titles and links if possible."
#             )
#             result = await loop.run_in_executor(None, lambda: llm.invoke(prompt))
#             response_text = result.content if hasattr(result, "content") else str(result)
#         except Exception as e:
#             response_text = f"❌ Error generating LLM response: {e}"

#     return {"result": response_text}







# backend.py - FastAPI Backend
import os
import asyncio
import requests
import xml.etree.ElementTree as ET
from typing import List
from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.base import Embeddings
from langchain_qdrant import Qdrant
from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance
from sentence_transformers import SentenceTransformer
from langchain_google_genai import ChatGoogleGenerativeAI
import PyPDF2
import docx

# Load environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv("GEMINI_API_KEY")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
QDRANT_URL = os.getenv("QDRANT_URL")

app = FastAPI()

VECTOR_DIM = 384  
COLLECTION_NAME = "uploaded_docs"
file_uploaded = False

# Keywords that indicate user wants to query the document
DOCUMENT_KEYWORDS = [
    "document", "pdf", "file", "uploaded", "from the document",
    "in the document", "according to the document", "explain from",
    "summarize the", "what does the document say", "extract from",
    "analyze the document", "document content", "based on the file",
    "from the pdf", "in the pdf", "what's in the", "content of the",
    "summary of the", "explain the document", "document summary"
]

# ------------------ Free Local Embeddings Wrapper ------------------
class LocalEmbeddings(Embeddings):
    def __init__(self):
        # Using a lightweight, fast model that runs locally
        self.model = SentenceTransformer('all-MiniLM-L6-v2')
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        embeddings = self.model.encode(texts)
        return embeddings.tolist()
    
    def embed_query(self, text: str) -> List[float]:
        embedding = self.model.encode([text])
        return embedding[0].tolist()

# Initialize embeddings
embeddings = LocalEmbeddings()
try:
    qdrant_client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=60)
    existing_collections = [c.name for c in qdrant_client.get_collections().collections]
    if COLLECTION_NAME not in existing_collections:
        qdrant_client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
        )
    vector_store = Qdrant(client=qdrant_client, collection_name=COLLECTION_NAME, embeddings=embeddings)
    print("✅ Qdrant Cloud is ready with local embeddings")
except Exception as e:
    vector_store = None
    print(f"❌ Qdrant unavailable: {e}")

# ------------------ Text Splitter ------------------
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

def get_gemini():
    return ChatGoogleGenerativeAI(
        model="gemini-2.0-flash-lite", 
        temperature=0.2, 
        api_key=GOOGLE_API_KEY,
        timeout=120
    )

def check_document_query(query: str) -> bool:
    """Check if the query is asking about the uploaded document."""
    query_lower = query.lower()
    return any(keyword in query_lower for keyword in DOCUMENT_KEYWORDS)

def search_arxiv_papers(query: str, max_results: int = 5):
    """Search ArXiv for research papers."""
    try:
        base_url = "http://export.arxiv.org/api/query"
        params = {
            'search_query': f'all:{query}',
            'start': 0,
            'max_results': max_results,
            'sortBy': 'relevance',
            'sortOrder': 'descending'
        }
        
        response = requests.get(base_url, params=params, timeout=15)
        if response.status_code != 200:
            return []
        
        root = ET.fromstring(response.content)
        papers = []
        for entry in root.findall('{http://www.w3.org/2005/Atom}entry'):
            title_elem = entry.find('{http://www.w3.org/2005/Atom}title')
            link_elem = entry.find('{http://www.w3.org/2005/Atom}id')
            summary_elem = entry.find('{http://www.w3.org/2005/Atom}summary')
            published_elem = entry.find('{http://www.w3.org/2005/Atom}published')
            authors = entry.findall('.//{http://www.w3.org/2005/Atom}author/{http://www.w3.org/2005/Atom}name')
            
            if title_elem is not None and link_elem is not None:
                title = title_elem.text.strip().replace('\n', ' ')
                link = link_elem.text.strip()
                summary = summary_elem.text.strip().replace('\n', ' ') if summary_elem is not None else "No summary"
                published = published_elem.text[:10] if published_elem is not None else "Unknown"
                author_names = [author.text for author in authors[:3]]
                
                papers.append({
                    'title': title,
                    'authors': ", ".join(author_names),
                    'link': link,
                    'summary': summary[:300] + "..." if len(summary) > 300 else summary,
                    'published': published,
                    'source': 'ArXiv'
                })
        
        return papers
    except Exception as e:
        print(f"ArXiv search error: {e}")
        return []

def search_pubmed_papers(query: str, max_results: int = 5):
    """Search PubMed for biomedical research papers."""
    try:
        # Search for PMIDs
        search_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
        search_params = {
            'db': 'pubmed',
            'term': query,
            'retmax': max_results,
            'retmode': 'json',
            'sort': 'relevance'
        }
        
        search_response = requests.get(search_url, params=search_params, timeout=10)
        search_data = search_response.json()
        
        pmids = search_data.get('esearchresult', {}).get('idlist', [])
        if not pmids:
            return []

        # Get paper details
        fetch_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
        fetch_params = {
            'db': 'pubmed',
            'id': ','.join(pmids),
            'retmode': 'xml'
        }
        
        fetch_response = requests.get(fetch_url, params=fetch_params, timeout=15)
        root = ET.fromstring(fetch_response.content)
        
        papers = []
        for article in root.findall('.//PubmedArticle'):
            title_elem = article.find('.//ArticleTitle')
            pmid_elem = article.find('.//PMID')
            abstract_elem = article.find('.//AbstractText')
            journal_elem = article.find('.//Journal/Title')
            date_elem = article.find('.//PubDate/Year')
            
            if title_elem is not None and pmid_elem is not None:
                title = ''.join(title_elem.itertext()).strip()
                pmid = pmid_elem.text
                abstract = ''.join(abstract_elem.itertext()).strip() if abstract_elem is not None else "No abstract"
                journal = journal_elem.text if journal_elem is not None else "Unknown Journal"
                year = date_elem.text if date_elem is not None else "Unknown"
                
                papers.append({
                    'title': title,
                    'authors': journal,  # Using journal as author info
                    'link': f'https://pubmed.ncbi.nlm.nih.gov/{pmid}/',
                    'summary': abstract[:300] + "..." if len(abstract) > 300 else abstract,
                    'published': year,
                    'source': 'PubMed'
                })
        
        return papers
    except Exception as e:
        print(f"PubMed search error: {e}")
        return []

def search_semantic_scholar(query: str, max_results: int = 5):
    """Search Semantic Scholar for research papers."""
    try:
        url = "https://api.semanticscholar.org/graph/v1/paper/search"
        params = {
            'query': query,
            'limit': max_results,
            'fields': 'title,authors,abstract,url,year,citationCount,journal'
        }
        
        response = requests.get(url, params=params, timeout=15)
        if response.status_code != 200:
            return []
        
        data = response.json()
        papers = []
        
        for paper in data.get('data', []):
            title = paper.get('title', 'No title')
            authors = [author.get('name', '') for author in paper.get('authors', [])[:3]]
            abstract = paper.get('abstract', 'No abstract available')
            url = paper.get('url', '')
            year = paper.get('year', 'Unknown')
            citations = paper.get('citationCount', 0)
            journal = paper.get('journal', {}).get('name', 'Unknown Journal') if paper.get('journal') else 'Unknown Journal'
            
            papers.append({
                'title': title,
                'authors': ", ".join(authors) if authors else journal,
                'link': url,
                'summary': abstract[:300] + "..." if len(abstract) > 300 else abstract,
                'published': str(year),
                'citations': citations,
                'source': 'Semantic Scholar'
            })

        return papers
    except Exception as e:
        print(f"Semantic Scholar search error: {e}")
        return []

async def get_research_papers(query: str):
    """Get research papers from multiple sources."""
    loop = asyncio.get_event_loop()
    
    # Run searches concurrently
    arxiv_task = loop.run_in_executor(None, search_arxiv_papers, query, 3)
    pubmed_task = loop.run_in_executor(None, search_pubmed_papers, query, 3)
    semantic_task = loop.run_in_executor(None, search_semantic_scholar, query, 4)
    
    try:
        arxiv_papers, pubmed_papers, semantic_papers = await asyncio.gather(
            arxiv_task, pubmed_task, semantic_task, return_exceptions=True
        )
        
        all_papers = []
        
        # Add papers from each source
        if isinstance(arxiv_papers, list):
            all_papers.extend(arxiv_papers)
        if isinstance(pubmed_papers, list):
            all_papers.extend(pubmed_papers)
        if isinstance(semantic_papers, list):
            all_papers.extend(semantic_papers)
        
        return all_papers[:10]  # Return top 10 papers
    
    except Exception as e:
        print(f"Error in research search: {e}")
        return []

def read_file(file: UploadFile) -> str:
    """Read PDF, DOCX, or TXT and return plain text."""
    try:
        if file.filename.endswith(".pdf"):
            file.file.seek(0)
            reader = PyPDF2.PdfReader(file.file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if not text.strip():
                raise ValueError("Could not extract text from PDF.")
            return text
        
        elif file.filename.endswith(".docx"):
            file.file.seek(0)
            doc = docx.Document(file.file)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        
        else:  # txt
            file.file.seek(0)
            return file.file.read().decode("utf-8")
    except Exception as e:
        raise ValueError(f"Error reading file: {str(e)}")

def process_text(text: str):
    """Split text into chunks and store in Qdrant."""
    if vector_store is None:
        raise Exception("Qdrant is not available")
    
    # Clear old collection content
    try:
        qdrant_client.delete_collection(collection_name=COLLECTION_NAME)
    except Exception:
        pass
    
    qdrant_client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
    )
    
    # Create new vector store instance
    vector_store_new = Qdrant(client=qdrant_client, collection_name=COLLECTION_NAME, embeddings=embeddings)
    
    doc = Document(page_content=text)
    chunks = text_splitter.split_documents([doc])
    vector_store_new.add_documents(chunks)
    
    return chunks

class Query(BaseModel):
    text: str

# ------------------ Routes ------------------
@app.post("/upload_file")
async def upload_file(file: UploadFile = File(...)):
    global file_uploaded, vector_store
    try:
        content = read_file(file)
        chunks = process_text(content)
        
        # Update vector_store reference
        vector_store = Qdrant(client=qdrant_client, collection_name=COLLECTION_NAME, embeddings=embeddings)
        file_uploaded = True
        
        return {
            "message": f"✅ File '{file.filename}' processed successfully with {len(chunks)} chunks.",
            "chunks": len(chunks)
        }
    except Exception as e:
        return {"message": f"❌ Error processing file: {str(e)}"}

@app.post("/chat")
async def chat(query: Query):
    user_text = query.text.strip()
    
    if not user_text:
        return {"result": "Please provide a valid query.", "mode": "error"}

    try:
        loop = asyncio.get_event_loop()
        is_document_query = check_document_query(user_text)

        async def call_llm(prompt_or_text):
            """Helper to call Gemini API with retry on rate limit."""
            llm = get_gemini()
            try:
                result = await asyncio.wait_for(
                    loop.run_in_executor(None, lambda: llm.invoke(prompt_or_text)),
                    timeout=60
                )
                return result
            except Exception as e:
                error_str = str(e)
                if "quota" in error_str.lower() and "retry_delay" in error_str.lower():
                    # Try to extract retry delay in seconds
                    import re, time
                    match = re.search(r"seconds: (\d+)", error_str)
                    delay = int(match.group(1)) if match else 60
                    print(f"⚠️ Rate limit hit. Retrying after {delay} seconds...")
                    await asyncio.sleep(delay)
                    # Retry once
                    result = await asyncio.wait_for(
                        loop.run_in_executor(None, lambda: llm.invoke(prompt_or_text)),
                        timeout=60
                    )
                    return result
                raise

        # ------------------ Document Mode ------------------
        if vector_store is not None and file_uploaded and is_document_query:
            relevant_docs = await loop.run_in_executor(
                None, lambda: vector_store.similarity_search(user_text, k=5)
            )
            
            if relevant_docs:
                context_text = "\n\n".join([doc.page_content for doc in relevant_docs])
                prompt = f"""Based on the following context from the uploaded document, answer the user's question.
                
Only use information from the provided context. If the answer is not in the context, say so.

Context:
{context_text}

Question: {user_text}

Answer:"""
                result = await call_llm(prompt)
                response_text = result.content if hasattr(result, "content") else str(result)
                return {"result": response_text, "mode": "document_search"}
            else:
                return {"result": "No relevant information found in the uploaded document.", "mode": "document_search"}

        # ------------------ Research Mode ------------------
        research_keywords = ["research papers", "papers on", "studies on", "academic papers", "publications", "scientific papers"]
        is_research_query = any(keyword in user_text.lower() for keyword in research_keywords)
        
        if is_research_query:
            search_query = user_text.lower()
            for keyword in research_keywords:
                search_query = search_query.replace(keyword, "").strip()
            search_query = search_query.replace("give", "").replace("find", "").replace("get", "").strip()
            
            papers = await get_research_papers(search_query)
            if not papers:
                return {"result": "No research papers found for this topic. Please try different keywords.", "mode": "research"}
            
            response = "📚 **Research Papers and Articles:**\n\n"
            for i, paper in enumerate(papers, 1):
                response += f"**{i}. {paper['title']}**\n"
                response += f"   📅 Published: {paper['published']}\n"
                response += f"   👥 Authors: {paper['authors']}\n"
                response += f"   🔗 Link: {paper['link']}\n"
                response += f"   📖 Summary: {paper['summary']}\n"
                response += f"   📊 Source: {paper['source']}\n"
                if 'citations' in paper:
                    response += f"   📈 Citations: {paper['citations']}\n"
                response += "\n" + "-"*50 + "\n\n"
            
            return {"result": response, "mode": "research", "papers_found": len(papers)}

        # ------------------ Normal LLM Mode ------------------
        result = await call_llm(user_text)
        response_text = result.content if hasattr(result, "content") else str(result)
        return {"result": response_text, "mode": "normal_llm"}

    except asyncio.TimeoutError:
        return {"result": "⏰ Request timed out. Please try with a shorter query.", "mode": "error"}
    except Exception as e:
        return {"result": f"❌ Error: {str(e)}", "mode": "error"}

@app.delete("/clear_documents")
async def clear_documents():
    global file_uploaded
    try:
        if vector_store is not None:
            qdrant_client.delete_collection(collection_name=COLLECTION_NAME)
            qdrant_client.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
            )
        file_uploaded = False
        return {"message": "✅ All documents cleared successfully."}
    except Exception as e:
        return {"message": f"❌ Error clearing documents: {str(e)}"}

@app.get("/status")
async def get_status():
    return {
        "status": "✅ API is running",
        "file_uploaded": file_uploaded,
        "qdrant_connected": vector_store is not None,
        "mode": "Intelligent Assistant (Document Q&A + Research + General Chat)"
    }


# backend.py - FastAPI Backend with Streaming + Memory
# import os
# import asyncio
# import requests
# import xml.etree.ElementTree as ET
# from typing import List, Optional, AsyncGenerator
# from fastapi import FastAPI, UploadFile, File
# from fastapi.responses import StreamingResponse, JSONResponse
# from pydantic import BaseModel
# from dotenv import load_dotenv
# from langchain.schema import Document
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.embeddings.base import Embeddings
# from langchain_qdrant import Qdrant
# from qdrant_client import QdrantClient
# from qdrant_client.models import VectorParams, Distance
# from sentence_transformers import SentenceTransformer
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.memory import ConversationBufferMemory
# import PyPDF2
# import docx

# # -------------------- Load environment --------------------
# load_dotenv()
# GOOGLE_API_KEY = os.getenv("GEMINI_API_KEY")
# QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
# QDRANT_URL = os.getenv("QDRANT_URL")

# # -------------------- App + state --------------------
# app = FastAPI(title="Intelligent Research & Document Assistant")

# VECTOR_DIM = 384
# COLLECTION_NAME = "uploaded_docs"
# file_uploaded: bool = False

# # Conversation memory
# memory = ConversationBufferMemory(return_messages=True)

# # -------------------- Local embeddings --------------------
# class LocalEmbeddings(Embeddings):
#     def __init__(self):
#         self.model = SentenceTransformer("all-MiniLM-L6-v2")

#     def embed_documents(self, texts: List[str]) -> List[List[float]]:
#         return self.model.encode(texts, show_progress_bar=False).tolist()

#     def embed_query(self, text: str) -> List[float]:
#         return self.model.encode([text], show_progress_bar=False)[0].tolist()

# embeddings = LocalEmbeddings()

# # -------------------- Qdrant init --------------------
# try:
#     qdrant_client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=60)
#     existing = [c.name for c in qdrant_client.get_collections().collections]
#     if COLLECTION_NAME not in existing:
#         qdrant_client.create_collection(
#             collection_name=COLLECTION_NAME,
#             vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
#         )
#     vector_store = Qdrant(client=qdrant_client, collection_name=COLLECTION_NAME, embeddings=embeddings)
#     print("✅ Qdrant ready.")
# except Exception as e:
#     qdrant_client, vector_store = None, None
#     print("❌ Qdrant unavailable:", e)

# # -------------------- Splitter --------------------
# text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

# # -------------------- Gemini wrapper --------------------
# def get_gemini():
#     return ChatGoogleGenerativeAI(
#         model="gemini-2.0-flash-lite",
#         temperature=0.2,
#         api_key=GOOGLE_API_KEY,
#         timeout=120
#     )

# # -------------------- Helpers --------------------
# DOCUMENT_KEYWORDS = ["document", "pdf", "file", "slide", "slides", "uploaded"]

# def check_document_query(query: str) -> bool:
#     return any(k in query.lower() for k in DOCUMENT_KEYWORDS)

# # -------------------- File reading --------------------
# def read_file(file: UploadFile) -> str:
#     if file.filename.lower().endswith(".pdf"):
#         reader = PyPDF2.PdfReader(file.file)
#         return "\n".join([p.extract_text() or "" for p in reader.pages])
#     elif file.filename.lower().endswith(".docx"):
#         doc = docx.Document(file.file)
#         return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
#     else:
#         raw = file.file.read()
#         return raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)

# def process_text_into_qdrant(text: str):
#     global vector_store, qdrant_client
#     if not qdrant_client:
#         raise Exception("Qdrant unavailable")
#     try:
#         qdrant_client.delete_collection(collection_name=COLLECTION_NAME)
#     except Exception:
#         pass
#     qdrant_client.create_collection(
#         collection_name=COLLECTION_NAME,
#         vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)
#     )
#     vector_store = Qdrant(client=qdrant_client, collection_name=COLLECTION_NAME, embeddings=embeddings)
#     doc = Document(page_content=text)
#     chunks = text_splitter.split_documents([doc])
#     vector_store.add_documents(chunks)
#     return chunks

# # -------------------- Research Search --------------------
# def search_arxiv(query: str, max_results=5): ...
# def search_pubmed(query: str, max_results=5): ...
# def search_semantic_scholar(query: str, max_results=5): ...

# # (Keep your earlier implementations here unchanged for brevity)

# async def get_research_papers(query: str):
#     loop = asyncio.get_event_loop()
#     return sum(await asyncio.gather(
#         loop.run_in_executor(None, search_arxiv, query, 3),
#         loop.run_in_executor(None, search_pubmed, query, 3),
#         loop.run_in_executor(None, search_semantic_scholar, query, 4),
#         return_exceptions=False
#     ), [])

# # -------------------- Pydantic --------------------
# class Query(BaseModel):
#     text: str

# # -------------------- Streaming Handler --------------------
# async def stream_response_with_memory(user_text: str) -> AsyncGenerator[str, None]:
#     llm = get_gemini()
#     memory.chat_memory.add_user_message(user_text)

#     lowered = user_text.lower()
#     research_keywords = ["research papers", "papers on", "studies on", "academic papers"]

#     try:
#         # -------- Document Mode --------
#         if check_document_query(user_text) and vector_store:
#             docs = await asyncio.get_event_loop().run_in_executor(
#                 None, lambda: vector_store.similarity_search(user_text, k=5)
#             )
#             if docs:
#                 context = "\n\n".join([d.page_content for d in docs])
#                 prompt = f"Answer ONLY using this context:\n{context}\n\nQ: {user_text}\nA:"
#                 final_text = ""
#                 for chunk in llm.stream(prompt):
#                     token = getattr(chunk, "content", str(chunk))
#                     if token:
#                         final_text += token
#                         yield token
#                 memory.chat_memory.add_ai_message(final_text)
#                 return
#             else:
#                 msg = "No relevant information found in the uploaded file."
#                 yield msg
#                 memory.chat_memory.add_ai_message(msg)
#                 return

#         # -------- Research Mode --------
#         if any(k in lowered for k in research_keywords):
#             query = user_text
#             for k in research_keywords:
#                 query = query.replace(k, "")
#             query = query.strip() or user_text
#             yield f"🔎 Searching for research papers on '{query}'...\n\n"
#             papers = await get_research_papers(query)
#             if not papers:
#                 msg = "No research papers found."
#                 yield msg
#                 memory.chat_memory.add_ai_message(msg)
#                 return
#             out = ""
#             for i, p in enumerate(papers, 1):
#                 entry = f"{i}. {p['title']} ({p['published']})\n   {p['link']}\n   {p['summary']}\n\n"
#                 out += entry
#                 yield entry
#                 await asyncio.sleep(0.02)
#             memory.chat_memory.add_ai_message(out)
#             return

#         # -------- Normal Mode --------
#         final_text = ""
#         for chunk in llm.stream(user_text):
#             token = getattr(chunk, "content", str(chunk))
#             if token:
#                 final_text += token
#                 yield token
#         memory.chat_memory.add_ai_message(final_text)

#     except Exception as e:
#         err = f"❌ Error: {e}"
#         yield err
#         memory.chat_memory.add_ai_message(err)

# # -------------------- Routes --------------------
# @app.post("/upload_file")
# async def upload_file(file: UploadFile = File(...)):
#     global file_uploaded
#     try:
#         content = read_file(file)
#         chunks = await asyncio.get_event_loop().run_in_executor(None, process_text_into_qdrant, content)
#         file_uploaded = True
#         return {"message": f"✅ '{file.filename}' processed with {len(chunks)} chunks."}
#     except Exception as e:
#         return JSONResponse(status_code=400, content={"message": str(e)})

# @app.post("/chat")
# async def chat_stream(query: Query):
#     if not query.text.strip():
#         return StreamingResponse(iter(["Please enter a query."]), media_type="text/plain")
#     return StreamingResponse(stream_response_with_memory(query.text.strip()), media_type="text/plain")

# @app.get("/status")
# async def status():
#     return {
#         "status": "✅ Running",
#         "file_uploaded": file_uploaded,
#         "messages_in_memory": len(memory.chat_memory.messages)
#     }
